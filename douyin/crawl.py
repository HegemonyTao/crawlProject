# -*- coding: UTF-8 –*-
'''
抖音爬虫，现可以进行个人信息的爬取以及某人下的视频(有水印+无水印)
作者：洪韬
时间：2020/3/2
'''
from urllib.parse import quote
from pyquery import PyQuery as pq
from docx.shared import Inches
from docx import Document
import requests
import time
import urllib
import random
import json
import re
import os


class DouYinCrawl:
    def __init__(self):
        self.__transDict = {'xe602': '1', 'xe603': '0', 'xe604': '3', 'xe605': '2', 'xe606': '4', 'xe607': '5',
                            'xe608': '6',
                            'xe609': '9', 'xe60a': '7', 'xe60b': '8', 'xe60c': '4', 'xe60d': '0', 'xe60e': '1',
                            'xe60f': '5',
                            'xe610': '2', 'xe611': '3', 'xe612': '6', 'xe613': '7', 'xe614': '8', 'xe615': '9',
                            'xe616': '0',
                            'xe617': '2', 'xe618': '1', 'xe619': '4', 'xe61a': '3', 'xe61b': '5', 'xe61c': '7',
                            'xe61d': '8',
                            'xe61e': '9', 'xe61f': '6'}
        self.__headers = {
            'upgrade-insecure-requests': '1',
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/80.0.3987.122 Safari/537.36',
            'sec-fetch-dest': 'document',
            'accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
            'sec-fetch-site': 'none',
            'sec-fetch-mode': 'navigate',
            'sec-fetch-user': '?1',
            'accept-language': 'zh-CN,zh;q=0.9',
            'cookie': '_ga=GA1.2.913540324.1582983429; _gid=GA1.2.1576801560.1582983429;'
        }

    def __safe_get(self, key, itemDict):
        try:
            return itemDict[key]
        except:
            return ''

    def __create_path(self, *path):
        finPath = ''
        for i in range(len(path)):
            finPath += path[i]
            if not os.path.exists(finPath):
                os.mkdir(finPath)
            finPath += '/'

    def __trans(self, transStr):
        finStr = transStr
        finStr = finStr.replace('&#', '').replace(';', '')
        for key in self.__transDict:
            finStr = finStr.replace(key, self.__transDict[key])
        return finStr

    def __add_param(self, document, boldWord, num):
        p = document.add_paragraph()
        p.add_run(boldWord).bold = True
        p.add_run(':' + num)

    # 得到某一个抖音号的信息
    def get_info(self, id, path):
        link = 'https://www.iesdouyin.com/share/user/' + str(id)
        response = requests.get(link, headers=self.__headers)
        html = pq(response.text)
        imgUrl = html('#pagelet-user-info > div.personal-card > div.info1 > span.author > img').attr('src')
        headers = {
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/80.0.3987.122 Safari/537.36'}
        tempRsp = requests.get(imgUrl, headers=headers)
        name = html('#pagelet-user-info > div.personal-card > div.info1 > p.nickname').text()
        self.__create_path(path, name)
        file = open(path + '/' + name + '/temp.png', 'wb')
        file.write(tempRsp.content)
        file.close()
        info = html('#pagelet-user-info > div.personal-card > div.info2 > div > span').text()
        signature = html('#pagelet-user-info > div.personal-card > div.info2 > p.signature').text()
        id = self.__trans(re.compile('<p class="shortid">(.*?)</p>').findall(response.text)[0])
        focus = self.__trans(re.compile('<span class="num">(.*?)</span>').findall(response.text)[0])
        follow = self.__trans(
            re.compile('<span class="follower block"><span class="num">(.*?)</span> </span>').findall(response.text)[0])
        liked = self.__trans(
            re.compile('<span class="liked-num block"><span class="num">(.*?)</span>').findall(response.text)[0])
        production = self.__trans(
            re.compile('<div class="user-tab active tab get-list" data-type="post">(.*?)</div>').findall(response.text)[
                0])
        like = self.__trans(
            re.compile('<div class="like-tab tab get-list" data-type="like">(.*?)</div>').findall(response.text)[0])
        id = pq(id).text().replace(' ', '').replace('抖音ID：', '')
        focus = pq(focus).text().replace(' ', '')
        follow = pq(follow).text().replace(' ', '').replace('粉丝', '')
        liked = pq(liked).text().replace(' ', '')
        production = pq(production).text().replace(' ', '').replace('作品', '')
        like = pq(like).text().replace(' ', '').replace('喜欢', '')
        document = Document()
        document.add_heading(name, 0)
        document.add_picture(path + '/' + name + '/temp.png', width=Inches(1.25))
        self.__add_param(document, '抖音ID', id)
        document.add_paragraph(info)
        document.add_paragraph(signature)
        self.__add_param(document, '关注', focus)
        self.__add_param(document, '粉丝', follow)
        self.__add_param(document, '赞', liked)
        self.__add_param(document, '作品', production)
        self.__add_param(document, '喜欢', like)
        document.save(path + '/' + name + '/info.docx')
        os.remove(path + '/' + name + '/temp.png')

    # 得到某个抖音号下的所有视频,water_mark用于标记是否有水印
    def get_videos(self, id, path,water_mark=True):
        response = requests.get('https://www.iesdouyin.com/share/user/' + str(id), headers=self.__headers)
        html = pq(response.text)
        name = html('#pagelet-user-info > div.personal-card > div.info1 > p.nickname').text()
        self.__create_path(path, name)
        production = self.__trans(
            re.compile('<div class="user-tab active tab get-list" data-type="post">(.*?)</div>').findall(response.text)[
                0])
        production = pq(production).text().replace(' ', '').replace('作品', '')
        cursor = 0
        print('共有'+str(production)+'个作品')
        session = requests.Session()
        url = "https://api.anoyi.com/api/signature/ies/" + str(id)
        headers = {
            'Connection': 'keep-alive',
            'Accept': '*/*',
            'Sec-Fetch-Dest': 'empty',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/80.0.3987.122 Safari/537.36',
            'Origin': 'https://anoyi.com',
            'Host': 'api.anoyi.com',
            'Sec-Fetch-Site': 'same-site',
            'Sec-Fetch-Mode': 'cors',
            'Accept-Language': 'zh-CN,zh;q=0.9'
        }
        response = session.get(url, headers=headers)
        jsonText = json.loads(response.text)
        _signature = jsonText['sign']
        dytk = jsonText['tk']
        if water_mark:
            print('正在爬取有水印视频')
        else:
            print('正在爬取无水印视频')
        page=1
        while True:
            time.sleep(2)
            print('正在爬取第' + str(page) + '页')
            page+=1
            url='https://api.anoyi.com/api/signature/ies/'+str(id)+'/post?tk='+dytk+'&max_cursor='+str(cursor)+'&s='+_signature
            response = session.get(url, headers=headers)
            jsonText = json.loads(response.text)
            cursor = jsonText['max_cursor']
            hasmore = jsonText['has_more']
            for item in jsonText['aweme_list']:
                title = item['desc']
                print('正在爬取视频"' + title + '"')
                url = item['video']['play_addr']['url_list'][0]  # 有水印视频
                if water_mark:
                    if os.path.exists(path + '/' + name + '/water_mark/' + title + '.mp4'):
                        print('视频已存在')
                        continue
                    url=url.replace('/play/','/playwm/')
                    headers = {
                        'authority': 'aweme.snssdk.com',
                        'sec-fetch-dest': 'video',
                        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/80.0.3987.163 Safari/537.36',
                        'accept': '*/*',
                        'sec-fetch-site': 'cross-site',
                        'sec-fetch-mode': 'no-cors',
                        'accept-language': 'zh-CN,zh;q=0.9',
                        'range': 'bytes=0-'
                    }
                    response = requests.get(url, headers=headers,allow_redirects=False)
                    url=response.headers['location']
                    headers = {
                        'Connection': 'keep-alive',
                        'Cache-Control': 'max-age=0',
                        'Upgrade-Insecure-Requests': '1',
                        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/80.0.3987.163 Safari/537.36',
                        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
                        'Accept-Language': 'zh-CN,zh;q=0.9'
                    }
                    response=requests.get(url,headers=headers)
                    self.__create_path('data',name,'water_mark')
                    file = open(path + '/' + name + '/water_mark/' + title + '.mp4', 'wb')
                    file.write(response.content)
                    file.close()
                else:
                    if os.path.exists(path + '/' + name + '/no_water_mark/' + title + '.mp4'):
                        print('视频已存在')
                        continue
                    url='https://api.anoyi.com/api/video/no-watermark?url='+quote(url)
                    headers = {
                        'Connection': 'keep-alive',
                        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/80.0.3987.163 Safari/537.36',
                        'Sec-Fetch-Dest': 'empty',
                        'Accept': '*/*',
                        'Origin': 'https://anoyi.com',
                        'Sec-Fetch-Site': 'same-site',
                        'Sec-Fetch-Mode': 'cors',
                        'Accept-Language': 'zh-CN,zh;q=0.9',
                        'If-None-Match': 'W/"148-8N1CAArBSSC2MpiZuqR7kZ2SBt4"'
                    }
                    rsp=requests.get(url,headers=headers)
                    headers = {
                        'Connection': 'keep-alive',
                        'Cache-Control': 'max-age=0',
                        'Upgrade-Insecure-Requests': '1',
                        'User-Agent': 'Mozilla/5.0 (iPhone; CPU iPhone OS 11_0 like Mac OS X) AppleWebKit/604.1.38 (KHTML, like Gecko) Version/11.0 Mobile/15A372 Safari/604.1',
                        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
                        'Accept-Language': 'zh-CN,zh;q=0.9'
                    }
                    response = requests.get(rsp.text, headers=headers)
                    self.__create_path('data',name,'no_water_mark')
                    file = open(path + '/' + name + '/no_water_mark/' + title + '.mp4', 'wb')
                    file.write(response.content)
                    file.close()
            if hasmore == False:
                break
if __name__ == '__main__':
    app = DouYinCrawl()
    # 得到个人信息
    # 第一个参数为个人id，第二个参数为存储的位置
    app.get_info('50548809333', 'data')
    # 得到个人视频
    # 第一个参数为个人id，第二个参数为存储的位置
    app.get_videos('50548809333', 'data',water_mark=True)
    app.get_videos('50548809333', 'data',water_mark=False)

